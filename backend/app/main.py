"""
FastAPI application for Candidate-Job Matching System

Core Architecture:
- Hybrid matching engine combining TF-IDF, semantic embeddings, skill scoring, and experience alignment
- RESTful API for job descriptions, candidates, and matching
- Support for file uploads (JSON, CSV)
- Caching of match results for performance
"""

import os
import json
import csv
import uuid
import re
from pathlib import Path
from typing import List, Optional
from datetime import datetime
import logging
from io import BytesIO

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import uvicorn
import openpyxl
from docx import Document

from .models import (
    JobDescription, CandidateProfile, MatchResult, JobMatchResults,
    DetailedMatchExplanation
)
from .database import Database
from .matching_engine import MatchingEngine

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Candidate-Job Matcher API",
    description="An intelligent matching system that ranks candidates for job positions",
    version="1.0.0"
)

# Add CORS middleware for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize components
database = Database()
matching_engine = MatchingEngine(use_embeddings=True)

# Global state
UPLOAD_DIR = Path("./uploads")
UPLOAD_DIR.mkdir(exist_ok=True)





def _sync_engine():
    """Synchronize database content with matching engine"""
    jobs = database.get_all_jobs()
    candidates = database.get_all_candidates()
    matching_engine.add_job_descriptions(jobs)
    matching_engine.add_candidates(candidates)
    logger.info(f"Engine synced: {len(jobs)} jobs, {len(candidates)} candidates")


# ============================================================================
# FILE PARSING HELPERS
# ============================================================================

def _parse_excel_file(content: bytes) -> List[dict]:
    """Parse Excel (.xlsx) file and return list of dictionaries
    Handles both structured tables and free-form data"""
    try:
        wb = openpyxl.load_workbook(BytesIO(content))
        ws = wb.active
        
        # Get all rows
        all_rows = list(ws.iter_rows(values_only=True))
        
        if not all_rows:
            logger.warning("Excel file is empty")
            return []
        
        # Check if first row looks like headers
        first_row = all_rows[0]
        logger.debug(f"First row of Excel: {first_row}")
        
        # Count non-empty cells that look like headers (text, not numbers)
        header_candidates = sum(1 for cell in first_row if cell and isinstance(cell, str) and len(str(cell)) < 50)
        has_headers = header_candidates > 1
        
        logger.debug(f"Excel header detection: {header_candidates} header-like cells, has_headers={has_headers}")
        
        data = []
        
        if has_headers and len(all_rows) > 1:
            # Structured table format
            headers = [str(cell).strip().lower().replace(' ', '_').replace('-', '_') if cell else f'col_{i}' 
                      for i, cell in enumerate(first_row)]
            
            logger.info(f"Excel headers found: {headers}")
            
            for row_idx, row in enumerate(all_rows[1:], start=2):
                if any(row):
                    row_dict = {}
                    for i, header in enumerate(headers):
                        value = row[i] if i < len(row) else None
                        if value is not None:
                            row_dict[header] = str(value).strip()
                    if row_dict:
                        logger.debug(f"Row {row_idx}: {row_dict}")
                        data.append(row_dict)
        else:
            # Free-form data - combine all non-empty cells into structured format
            combined_text = []
            for row in all_rows:
                for cell in row:
                    if cell:
                        combined_text.append(str(cell).strip())
            
            if combined_text:
                # Create a single entry from all content
                full_text = ' '.join(combined_text)
                logger.debug(f"Free-form Excel data combined into: {full_text[:100]}...")
                data.append({
                    'content': full_text,
                    'title': combined_text[0] if combined_text else 'Unnamed',
                    'description': ' '.join(combined_text[1:]) if len(combined_text) > 1 else ''
                })
        
        logger.info(f"Parsed Excel file: {len(data)} entries")
        return data
    except Exception as e:
        logger.error(f"Error parsing Excel file: {e}", exc_info=True)
        raise ValueError(f"Failed to parse Excel file: {str(e)}")


def _parse_word_file(content: bytes) -> List[dict]:
    """Parse Word (.docx) file and return list of dictionaries
    Handles both tables and narrative text"""
    try:
        doc = Document(BytesIO(content))
        data = []
        
        # First try to extract from tables (for structured candidate/job data)
        if doc.tables:
            for table in doc.tables:
                headers = [cell.text.strip().lower().replace(' ', '_') for cell in table.rows[0].cells]
                
                for row in table.rows[1:]:
                    row_dict = {}
                    for i, header in enumerate(headers):
                        if i < len(row.cells):
                            text = row.cells[i].text.strip()
                            if text:
                                row_dict[header] = text
                    
                    if row_dict:
                        data.append(row_dict)
        
        # If no tables or minimal data, extract from paragraphs (for narrative documents)
        if not data and doc.paragraphs:
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                # Group paragraphs into logical entries
                full_text = ' '.join(paragraphs)
                
                # Try to split by common section headers
                sections = {
                    'content': full_text,
                    'title': paragraphs[0] if paragraphs else 'Unnamed',
                    'description': ' '.join(paragraphs[1:]) if len(paragraphs) > 1 else ''
                }
                data.append(sections)
        
        if not data:
            raise ValueError("No extractable data found in Word document")
        
        return data
    except Exception as e:
        logger.error(f"Error parsing Word file: {e}")
        raise ValueError(f"Failed to parse Word file: {str(e)}")




def _extract_skills_from_text(text: str) -> List[str]:
    """Extract skills from job description text by looking for 'Skills Required' or similar sections"""
    if not text:
        logger.warning("[SKILL EXTRACTION] No text provided")
        return []
    
    text = str(text).strip()
    skills = []
    
    logger.warning(f"[SKILL EXTRACTION] Attempting to extract from text (length: {len(text)} chars)")
    
    # Pattern 1: Find "Skills Required" section and capture everything after it
    # Look for the "Skills Required" text, then get everything until the next section or end
    matches = list(re.finditer(r'(?:skills\s+(?:required|needed)|required\s+skills)', text, re.IGNORECASE))
    
    if matches:
        logger.warning(f"[SKILL EXTRACTION] Found {len(matches)} 'Skills Required' section(s)")
        match_start = matches[0].end()  # Position right after "Skills Required"
        
        # Find where the skills text starts and ends
        # Skills typically start immediately after the section header
        remaining_text = text[match_start:]
        
        # Remove leading colons, spaces, and newlines
        remaining_text = re.sub(r'^[\s:]*\n*', '', remaining_text)
        
        logger.warning(f"[SKILL EXTRACTION] Text after 'Skills Required': {remaining_text[:100]}...")
        
        # Find the end of the skills section (either double newline, or next major section)
        end_match = re.search(r'\n\n|(?:\n[A-Z][a-zA-Z\s]*:)|$', remaining_text)
        if end_match:
            skills_text = remaining_text[:end_match.start()].strip()
        else:
            skills_text = remaining_text.strip()
        
        logger.warning(f"[SKILL EXTRACTION] Extracted skills text: {skills_text[:150]}...")
        
        # Split by comma, semicolon, or newline
        if skills_text:
            skills = [s.strip() for s in re.split(r'[,;\n•-]', skills_text) if s.strip()]
            logger.warning(f"[SKILL EXTRACTION] Found {len(skills)} skills: {skills[:5]}...")
            return skills
    
    logger.warning(f"[SKILL EXTRACTION] No 'Skills Required' section found")
    
    # Fallback: Look for any line with multiple comma-separated values
    lines = text.split('\n')
    for line_num, line in enumerate(lines):
        comma_count = line.count(',')
        if comma_count >= 3 and len(line) > 30:
            logger.warning(f"[SKILL EXTRACTION] Found comma-separated line {line_num} with {comma_count} commas: {line[:150]}...")
            skills = [s.strip() for s in line.split(',') if s.strip()]
            logger.warning(f"[SKILL EXTRACTION] Extracted {len(skills)} skills from comma list: {skills[:5]}...")
            return skills
    
    logger.warning(f"[SKILL EXTRACTION] No skills found - extracted list is empty")
    return []

def _extract_job_from_entry(entry: dict) -> dict:
    """Extract job description data from a flexible entry format"""
    job = {}
    
    # Normalize keys first (lowercase, underscores)
    normalized_entry = {}
    for key, value in entry.items():
        if key:
            norm_key = key.strip().lower().replace(' ', '_').replace('-', '_')
            normalized_entry[norm_key] = value
    
    # Log all available normalized fields for first job
    if not hasattr(_extract_job_from_entry, 'logged_fields'):
        logger.warning(f"AVAILABLE FIELDS IN JOB DATA: {list(normalized_entry.keys())}")
        _extract_job_from_entry.logged_fields = True
    
    # Map various field names to standard job fields
    field_mappings = {
        'title': ['title', 'job_title', 'position', 'role', 'job_position', 'job_name'],
        'description': ['description', 'job_description', 'details', 'content', 'job_details', 'about_role', 'role_description'],
        'required_skills': ['required_skills', 'must_have_skills', 'must_have', 'core_skills', 'required_skills_list', 'required', 'skills_required', 'required_competencies'],
        'optional_skills': ['optional_skills', 'nice_to_have_skills', 'nice_to_have', 'preferred_skills', 'preferred', 'good_to_have'],
        'company': ['company', 'organization', 'employer', 'company_name', 'hiring_company'],
        'location': ['location', 'job_location', 'office', 'place', 'work_location', 'office_location'],
        'years_required': ['years_required', 'experience_required', 'experience', 'min_years', 'min_experience', 'years_of_experience_required', 'experience_level_years'],
        'experience_level': ['experience_level', 'level', 'seniority', 'seniority_level', 'job_level', 'career_level'],
        'salary_range': ['salary_range', 'salary', 'compensation', 'pay', 'salary_band', 'salary_range_usd']
    }
    
    # Try to map fields from entry
    for target_field, source_fields in field_mappings.items():
        for source_field in source_fields:
            if source_field in normalized_entry and normalized_entry[source_field]:
                job[target_field] = normalized_entry[source_field]
                logger.debug(f"Mapped job field {source_field} to {target_field}: {normalized_entry[source_field]}")
                break
    
    # DEBUG: Log what happened with skills specifically
    job_title = job.get('title', 'Unknown')
    if 'required_skills' in job and job['required_skills']:
        logger.warning(f"[JOB REQUIRED SKILLS FOUND] {job_title}: {len(job['required_skills'])} skills")
    else:
        logger.warning(f"[NO JOB REQUIRED SKILLS IN STRUCTURED FIELDS] {job_title}: Will try extracting from description...")
        # Try extracting from job description text
        if 'description' in job and job['description']:
            extracted_skills = _extract_skills_from_text(job['description'])
            if extracted_skills:
                job['required_skills'] = extracted_skills
                logger.warning(f"[SUCCESS] Extracted {len(extracted_skills)} skills from job description")
    
    # Convert skills to list if string
    if 'required_skills' in job and isinstance(job['required_skills'], str):
        job['required_skills'] = _normalize_skills(job['required_skills'])
    
    if 'optional_skills' in job and isinstance(job['optional_skills'], str):
        job['optional_skills'] = _normalize_skills(job['optional_skills'])
    
    # Ensure minimum required fields
    if 'title' not in job or not job['title']:
        job['title'] = normalized_entry.get('title', 'Untitled Job')
    
    if 'description' not in job or not job['description']:
        job['description'] = normalized_entry.get('description', normalized_entry.get('content', ''))
    
    # Ensure skills are lists
    if 'required_skills' not in job or not job['required_skills']:
        job['required_skills'] = []
    
    if 'optional_skills' not in job or not job['optional_skills']:
        job['optional_skills'] = []
    
    if 'id' not in job:
        job['id'] = f"jd-{uuid.uuid4().hex[:8]}"
    
    return job


def _extract_candidate_from_entry(entry: dict) -> dict:
    """Extract candidate profile data from a flexible entry format"""
    candidate = {}
    
    # Normalize keys first (lowercase, underscores)
    normalized_entry = {}
    for key, value in entry.items():
        if key:
            norm_key = key.strip().lower().replace(' ', '_').replace('-', '_')
            normalized_entry[norm_key] = value
    
    # Log all available normalized fields for first candidate
    if not hasattr(_extract_candidate_from_entry, 'logged_fields'):
        logger.warning(f"AVAILABLE FIELDS IN CSV: {list(normalized_entry.keys())}")
        _extract_candidate_from_entry.logged_fields = True
    
    # Map various field names to standard candidate fields
    field_mappings = {
        'name': ['name', 'full_name', 'candidate_name', 'applicant_name', 'candidate', 'fullname'],
        'summary': ['summary', 'bio', 'about', 'profile', 'description', 'content', 'professional_summary', 'overview', 'parsed_summary'],
        'skills': ['skills', 'skill', 'technical_skills', 'competencies', 'expertise', 'tecnical_skils', 'tech_skills', 'key_skills', 'main_skills', 'abilities', 'parsed_skills', 'all_skills'],
        'experience_years': ['experience_years', 'years_experience', 'years', 'experience', 'exp_years', 'total_experience', 'years_of_experience', 'yrs_experience', 'parsed_metadata_calculated_years_experience'],
        'email': ['email', 'email_address', 'email_id', 'work_email'],
        'phone': ['phone', 'phone_number', 'contact_phone', 'mobile', 'mobile_number'],
        'location': ['location', 'city', 'place', 'based_in', 'current_location', 'region'],
        'education': ['education', 'degree', 'university', 'qualification', 'educational_qualification', 'academic', 'parsed_metadata_education', 'education_status'],
        'certifications': ['certifications', 'cert', 'credentials', 'certificates', 'certs', 'qualifications'],
        'previous_roles': ['previous_roles', 'work_experience', 'experience', 'past_roles', 'employment', 'job_titles', 'work_history', 'parsed_work_experience', 'current_title', 'recent_experience_type'],
        'linkedin_url': ['linkedin', 'linkedin_url', 'linkedin_profile', 'linkedin_id'],
        'github_url': ['github', 'github_url', 'github_profile', 'github_id', 'github_username']
    }
    
    # Additional skill field sources to combine
    skill_field_alternates = ['programming_languages', 'backend_frameworks', 'frontend_technologies', 'mobile_technologies']
    
    # Try to map fields from entry
    for target_field, source_fields in field_mappings.items():
        for source_field in source_fields:
            if source_field in normalized_entry and normalized_entry[source_field]:
                candidate[target_field] = normalized_entry[source_field]
                logger.debug(f"Mapped {source_field} to {target_field}: {normalized_entry[source_field]}")
                break
    
    # DEBUG: Log what happened with skills specifically
    candidate_name = candidate.get('name', 'Unknown')
    if 'skills' in candidate:
        if isinstance(candidate['skills'], list):
            logger.warning(f"[SKILLS FOUND AS LIST] {candidate_name}: {len(candidate['skills'])} skills - {candidate['skills'][:5]}...")
        else:
            logger.warning(f"[SKILLS FOUND AS STRING] {candidate_name}: {candidate['skills'][:100]}...")
    else:
        logger.warning(f"[NO SKILLS] {candidate_name}: Checking alternate fields...")
        logger.warning(f"  - parsed_skills: {normalized_entry.get('parsed_skills', 'NOT FOUND')[:100] if normalized_entry.get('parsed_skills') else 'NOT FOUND'}")
        logger.warning(f"  - programming_languages: {normalized_entry.get('programming_languages', 'NOT FOUND')}")
        logger.warning(f"  - backend_frameworks: {normalized_entry.get('backend_frameworks', 'NOT FOUND')}")
        logger.warning(f"  - frontend_technologies: {normalized_entry.get('frontend_technologies', 'NOT FOUND')}")
        logger.warning(f"  - mobile_technologies: {normalized_entry.get('mobile_technologies', 'NOT FOUND')}")
    
    # Convert skills to list if it's a string (CRITICAL: parsed_skills comes as string)
    if 'skills' in candidate and isinstance(candidate['skills'], str):
        candidate['skills'] = _normalize_skills(candidate['skills'])
        logger.warning(f"[CONVERTED SKILLS TO LIST] {candidate_name}: {len(candidate['skills'])} skills after normalization")
    
    # If skills not found, try combining alternate skill fields
    if 'skills' not in candidate or not candidate['skills']:
        combined_skills = []
        for skill_field in skill_field_alternates:
            if skill_field in normalized_entry and normalized_entry[skill_field]:
                field_value = normalized_entry[skill_field]
                if isinstance(field_value, str):
                    combined_skills.extend(_normalize_skills(field_value))
                elif isinstance(field_value, list):
                    combined_skills.extend(field_value)
        
        if combined_skills:
            candidate['skills'] = combined_skills
            logger.warning(f"[COMBINED SKILLS] {candidate_name}: {combined_skills}")
    
    # Convert skills to list if it's a string
    if 'skills' in candidate and isinstance(candidate['skills'], str):
        candidate['skills'] = _normalize_skills(candidate['skills'])
    
    # Ensure minimum required fields
    if 'name' not in candidate or not candidate['name']:
        candidate['name'] = normalized_entry.get('name', 'Unnamed Candidate')
    
    # If summary is missing, construct it from available fields
    if 'summary' not in candidate or not candidate['summary']:
        summary_parts = []
        if candidate.get('experience_years'):
            summary_parts.append(f"{candidate['experience_years']} years of experience")
        if candidate.get('education'):
            summary_parts.append(f"Education: {candidate['education']}")
        if candidate.get('profile'):
            summary_parts.append(candidate['profile'])
        
        # If still nothing, use skills or a generic message
        if summary_parts:
            candidate['summary'] = '. '.join(summary_parts)
        else:
            candidate['summary'] = f"Professional profile for {candidate.get('name', 'candidate')}"
    
    # Ensure skills is always a list
    if 'skills' not in candidate or not candidate['skills']:
        candidate['skills'] = []
    
    # Ensure certifications is always a list
    if 'certifications' not in candidate or not candidate['certifications']:
        candidate['certifications'] = []
    
    # Ensure previous_roles is always a list
    if 'previous_roles' not in candidate or not candidate['previous_roles']:
        candidate['previous_roles'] = []
    
    if 'id' not in candidate:
        candidate['id'] = f"cand-{uuid.uuid4().hex[:8]}"
    
    return candidate


def _normalize_skills(skills_str: str) -> List[str]:
    """Convert skills string to list, handling various separators and formats"""
    if not skills_str:
        return []
    
    if isinstance(skills_str, list):
        return skills_str
    
    # Convert to string just in case
    skills_str = str(skills_str).strip()
    
    if not skills_str:
        return []
    
    # Try to parse as JSON array first
    try:
        import json
        parsed = json.loads(skills_str)
        if isinstance(parsed, list):
            return [str(s).strip() for s in parsed if s]
    except:
        pass
    
    # Try different separators in order of likelihood
    separators = [',', ';', '\\n', '|', '\n']
    for sep in separators:
        if sep in skills_str:
            skills = [s.strip() for s in skills_str.split(sep) if s.strip()]
            if len(skills) > 1:  # Only use this separator if it actually splits into multiple items
                logger.debug(f"Split skills using '{sep}': {skills}")
                return skills
    
    # If no separator found, return as single-item list
    return [skills_str] if skills_str else []


# ============================================================================
# STARTUP AND SHUTDOWN
# ============================================================================

@app.on_event("startup")
async def startup_event():
    """Initialize the application"""
    logger.info("Starting Candidate Matcher Application")
    _sync_engine()
    logger.info("Application ready")


# ============================================================================
# HEALTH AND STATUS ENDPOINTS
# ============================================================================

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


@app.get("/stats")
async def get_stats():
    """Get system statistics"""
    jobs = database.get_all_jobs()
    candidates = database.get_all_candidates()
    return {
        "total_jobs": len(jobs),
        "total_candidates": len(candidates),
        "database_type": "SQLite",
        "matching_engine": "Hybrid (TF-IDF + Semantic + Skills + Experience)"
    }


# ============================================================================
# JOB DESCRIPTION ENDPOINTS
# ============================================================================

@app.post("/jobs", response_model=dict)
async def create_job(job: JobDescription):
    """
    Create a new job description
    
    Request body should include:
    - title: Job title
    - description: Full job description
    - required_skills: List of required skills
    - optional_skills: List of optional skills
    - experience_level: junior, mid, or senior
    - years_required: Years of experience required
    """
    if not job.id:
        job.id = f"jd-{uuid.uuid4().hex[:8]}"
    
    job.created_at = datetime.utcnow()
    job_dict = job.dict()
    
    if database.add_job(job_dict):
        matching_engine.add_job_descriptions([job_dict])
        logger.info(f"Created job: {job.id}")
        return {"id": job.id, "message": "Job created successfully"}
    else:
        raise HTTPException(status_code=500, detail="Failed to create job")


@app.get("/jobs")
async def list_jobs():
    """List all job descriptions"""
    jobs = database.get_all_jobs()
    return {"total": len(jobs), "jobs": jobs}


@app.get("/jobs/{job_id}")
async def get_job(job_id: str):
    """Get a specific job description"""
    job = database.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@app.post("/jobs/upload")
async def upload_jobs(file: UploadFile = File(...)):
    """
    Upload job descriptions from ANY file type (JSON, CSV, Excel, Word, Text, etc.)
    
    The system will intelligently extract meaningful data from:
    - JSON: Structured job objects
    - CSV: Comma-separated values with headers
    - Excel (.xlsx): Structured tables or free-form data
    - Word (.docx): Tables or narrative text
    - Text files: Job descriptions as paragraphs
    
    Field names are flexible and will be automatically mapped.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    content = await file.read()
    jobs_added = 0
    errors = []
    
    try:
        jobs_data = []
        
        if file.filename.endswith('.json'):
            data = json.loads(content)
            jobs_data = data if isinstance(data, list) else [data]
        
        elif file.filename.endswith('.csv'):
            text = content.decode('utf-8')
            reader = csv.DictReader(text.splitlines())
            jobs_data = [dict(row) for row in reader]
        
        elif file.filename.endswith('.xlsx'):
            jobs_data = _parse_excel_file(content)
        
        elif file.filename.endswith('.docx'):
            jobs_data = _parse_word_file(content)
        
        elif file.filename.endswith('.txt'):
            # Parse as plain text job description
            text = content.decode('utf-8')
            jobs_data = [{'title': 'Job from Text', 'description': text}]
        
        else:
            # Default: try to parse as JSON, then as text
            try:
                data = json.loads(content)
                jobs_data = data if isinstance(data, list) else [data]
            except:
                text = content.decode('utf-8')
                jobs_data = [{'title': file.filename, 'description': text}]
        
        # Process and add jobs
        for i, entry in enumerate(jobs_data):
            try:
                # Extract job data intelligently from flexible format
                job_data = _extract_job_from_entry(entry)
                logger.warning(f"[UPLOAD] Extracted job {i+1}: title={job_data.get('title')}, has_required_skills={bool(job_data.get('required_skills'))}, count={len(job_data.get('required_skills', []))}")
                
                # Validate minimum required fields
                if not job_data.get('title') or not job_data.get('description'):
                    errors.append(f"Entry {i+1}: Missing title or description")
                    logger.warning(f"Entry {i+1} rejected: title={job_data.get('title')}, description={job_data.get('description')}")
                    continue
                
                # Parse skills if they're strings
                if isinstance(job_data.get('required_skills'), str):
                    job_data['required_skills'] = _normalize_skills(job_data['required_skills'])
                    logger.warning(f"[UPLOAD] Converted string skills to list: {len(job_data['required_skills'])} skills")
                elif not job_data.get('required_skills'):
                    job_data['required_skills'] = []
                
                if isinstance(job_data.get('optional_skills'), str):
                    job_data['optional_skills'] = _normalize_skills(job_data['optional_skills'])
                elif not job_data.get('optional_skills'):
                    job_data['optional_skills'] = []
                
                # Parse years_required
                try:
                    job_data['years_required'] = int(job_data.get('years_required', 0))
                except (ValueError, TypeError):
                    job_data['years_required'] = 0
                
                # Final validation: ensure skills are set
                logger.warning(f"[UPLOAD PRE-SAVE] Job: {job_data.get('title')} | Required Skills Count: {len(job_data.get('required_skills', []))} | Optional Skills Count: {len(job_data.get('optional_skills', []))}")
                if job_data.get('required_skills'):
                    logger.warning(f"[UPLOAD PRE-SAVE] Sample required skills: {job_data['required_skills'][:3]}")
                
                if database.add_job(job_data):
                    matching_engine.add_job_descriptions([job_data])
                    jobs_added += 1
                    logger.info(f"Added job: {job_data.get('title')} | Required Skills: {job_data.get('required_skills')} | Optional Skills: {job_data.get('optional_skills')} | Experience: {job_data.get('years_required')}y")
            except Exception as e:
                errors.append(f"Entry {i+1}: {str(e)}")
                logger.error(f"Error processing entry {i+1}: {e}", exc_info=True)
        
        return {
            "message": f"Uploaded {jobs_added} jobs",
            "count": jobs_added,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        logger.error(f"Error uploading jobs: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")


# ============================================================================
# CANDIDATE PROFILE ENDPOINTS
# ============================================================================

@app.post("/candidates", response_model=dict)
async def create_candidate(candidate: CandidateProfile):
    """
    Create a new candidate profile
    
    Request body should include:
    - name: Candidate name
    - summary: Professional summary
    - skills: List of skills
    - experience_years: Years of experience
    """
    if not candidate.id:
        candidate.id = f"cand-{uuid.uuid4().hex[:8]}"
    
    candidate.created_at = datetime.utcnow()
    candidate_dict = candidate.dict()
    
    if database.add_candidate(candidate_dict):
        matching_engine.add_candidates([candidate_dict])
        logger.info(f"Created candidate: {candidate.id}")
        return {"id": candidate.id, "message": "Candidate created successfully"}
    else:
        raise HTTPException(status_code=500, detail="Failed to create candidate")


@app.get("/candidates")
async def list_candidates():
    """List all candidate profiles"""
    candidates = database.get_all_candidates()
    return {"total": len(candidates), "candidates": candidates}


@app.get("/candidates/{candidate_id}")
async def get_candidate(candidate_id: str):
    """Get a specific candidate profile"""
    candidate = database.get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    return candidate


@app.post("/candidates/upload")
async def upload_candidates(file: UploadFile = File(...)):
    """
    Upload candidate profiles from ANY file type (JSON, CSV, Excel, Word, Text, etc.)
    
    The system will intelligently extract meaningful data from:
    - JSON: Structured candidate objects
    - CSV: Comma-separated values with headers
    - Excel (.xlsx): Structured tables or free-form data
    - Word (.docx): Tables or narrative text (paragraphs become profile content)
    - Text files: Treated as candidate summary
    
    Field names are flexible and will be automatically mapped to:
    name, summary, skills, experience_years, email, phone, education, certifications, etc.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    content = await file.read()
    candidates_added = 0
    errors = []
    
    try:
        candidates_data = []
        
        if file.filename.endswith('.json'):
            data = json.loads(content)
            candidates_data = data if isinstance(data, list) else [data]
        
        elif file.filename.endswith('.csv'):
            text = content.decode('utf-8')
            reader = csv.DictReader(text.splitlines())
            candidates_data = [dict(row) for row in reader]
        
        elif file.filename.endswith('.xlsx'):
            candidates_data = _parse_excel_file(content)
        
        elif file.filename.endswith('.docx'):
            candidates_data = _parse_word_file(content)
        
        elif file.filename.endswith('.txt'):
            # Parse as plain text candidate summary
            text = content.decode('utf-8')
            candidates_data = [{'name': 'Candidate from Text', 'summary': text}]
        
        else:
            # Default: try to parse as JSON, then as text
            try:
                data = json.loads(content)
                candidates_data = data if isinstance(data, list) else [data]
            except:
                text = content.decode('utf-8')
                candidates_data = [{'name': file.filename, 'summary': text}]
        
        logger.info(f"Parsed {len(candidates_data)} entries from {file.filename}")
        if candidates_data:
            logger.debug(f"First entry sample keys: {list(candidates_data[0].keys())}")
        
        # Process and add candidates
        for i, entry in enumerate(candidates_data):
            try:
                # Log raw entry before extraction
                logger.debug(f"Entry {i+1} raw skills fields - parsed_skills: {entry.get('parsed_skills')}, programming_languages: {entry.get('programming_languages')}, backend_frameworks: {entry.get('backend_frameworks')}, frontend_technologies: {entry.get('frontend_technologies')}")
                
                # Extract candidate data intelligently from flexible format
                candidate_data = _extract_candidate_from_entry(entry)
                logger.debug(f"Extracted candidate {i+1}: name={candidate_data.get('name')}, skills={candidate_data.get('skills')}, skills_count={len(candidate_data.get('skills', []))}")
                
                # Validate minimum required fields
                if not candidate_data.get('name') or not candidate_data.get('summary'):
                    errors.append(f"Entry {i+1}: Missing name or summary")
                    logger.warning(f"Entry {i+1} rejected: name={candidate_data.get('name')}, summary={bool(candidate_data.get('summary'))}")
                    continue
                
                # Parse skills if they're strings
                if isinstance(candidate_data.get('skills'), str):
                    candidate_data['skills'] = _normalize_skills(candidate_data['skills'])
                elif not candidate_data.get('skills'):
                    candidate_data['skills'] = []
                
                # Log after skill parsing
                logger.info(f"Candidate {i+1} ({candidate_data.get('name')}): skills={candidate_data.get('skills')}")
                
                # Parse certifications if they're strings
                if isinstance(candidate_data.get('certifications'), str):
                    candidate_data['certifications'] = _normalize_skills(candidate_data['certifications'])
                elif not candidate_data.get('certifications'):
                    candidate_data['certifications'] = []
                
                # Parse experience_years
                try:
                    candidate_data['experience_years'] = int(candidate_data.get('experience_years', 0))
                except (ValueError, TypeError):
                    candidate_data['experience_years'] = 0
                
                if database.add_candidate(candidate_data):
                    matching_engine.add_candidates([candidate_data])
                    candidates_added += 1
                    logger.info(f"Added candidate: {candidate_data.get('name')} | Skills: {candidate_data.get('skills')} | Experience: {candidate_data.get('experience_years')}y")
            except Exception as e:
                errors.append(f"Entry {i+1}: {str(e)}")
                logger.error(f"Error processing entry {i+1}: {e}", exc_info=True)
        
        return {
            "message": f"Uploaded {candidates_added} candidates",
            "count": candidates_added,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        logger.error(f"Error uploading candidates: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")


# ============================================================================
# MATCHING ENDPOINTS
# ============================================================================

@app.get("/match/{jd_id}", response_model=JobMatchResults)
async def get_matches(jd_id: str, limit: int = 10):
    """
    Get ranked candidates for a specific job description
    
    Returns:
    - List of candidates ranked by match score
    - Each candidate includes overall score and explanation
    - Results sorted by score (highest first)
    """
    job = database.get_job(jd_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    try:
        matches = matching_engine.match_candidates(jd_id)
        
        # Limit results
        matches = matches[:limit]
        
        results = []
        for rank, match in enumerate(matches, 1):
            result = MatchResult(
                candidate_id=match.candidate_id,
                candidate_name=database.get_candidate(match.candidate_id).get('name'),
                rank=rank,
                overall_score=match.overall_score,
                keyword_score=match.keyword_score,
                semantic_score=match.semantic_score,
                skill_score=match.skill_score,
                experience_score=match.experience_score,
                explanation_summary=match.explanation_summary,
                detailed_explanation=match.detailed_explanation
            )
            results.append(result)
            
            # Cache the result
            database.save_match_result({
                'job_id': jd_id,
                'candidate_id': match.candidate_id,
                'overall_score': match.overall_score,
                'keyword_score': match.keyword_score,
                'semantic_score': match.semantic_score,
                'skill_score': match.skill_score,
                'experience_score': match.experience_score,
                'explanation_summary': match.explanation_summary,
                'detailed_explanation': match.detailed_explanation
            })
        
        return JobMatchResults(
            job_id=jd_id,
            job_title=job.get('title'),
            total_candidates=len(matches),
            timestamp=datetime.utcnow(),
            results=results
        )
    
    except Exception as e:
        logger.error(f"Error matching candidates: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/match/{jd_id}/{candidate_id}")
async def get_match_detail(jd_id: str, candidate_id: str):
    """
    Get detailed matching explanation for a specific candidate-job pair
    
    Returns:
    - Detailed breakdown of matching scores
    - Skill match analysis
    - Experience alignment
    - Technology/keyword matching
    - Human-readable explanation
    """
    job = database.get_job(jd_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    candidate = database.get_candidate(candidate_id)
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    try:
        match = matching_engine.get_match_details(jd_id, candidate_id)
        
        return {
            "job_id": jd_id,
            "job_title": job.get('title'),
            "candidate_id": candidate_id,
            "candidate_name": candidate.get('name'),
            "overall_score": match.overall_score,
            "explanation_summary": match.explanation_summary,
            "detailed_explanation": match.detailed_explanation,
            "score_breakdown": {
                "skill_score": match.skill_score,
                "experience_score": match.experience_score,
                "keyword_score": match.keyword_score,
                "semantic_score": match.semantic_score
            },
            "job_requirements": {
                "title": job.get('title'),
                "required_skills": job.get('required_skills', []),
                "optional_skills": job.get('optional_skills', []),
                "years_required": job.get('years_required'),
                "company": job.get('company'),
                "location": job.get('location')
            },
            "candidate_profile": {
                "name": candidate.get('name'),
                "summary": candidate.get('summary'),
                "skills": candidate.get('skills', []),
                "experience_years": candidate.get('experience_years'),
                "location": candidate.get('location'),
                "certifications": candidate.get('certifications', [])
            }
        }
    
    except Exception as e:
        logger.error(f"Error getting match details: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/match/{jd_id}/summary")
async def get_match_summary(jd_id: str):
    """Get summary statistics for all matches for a job"""
    job = database.get_job(jd_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    try:
        matches = matching_engine.match_candidates(jd_id)
        
        scores = [m.overall_score for m in matches]
        
        return {
            "job_id": jd_id,
            "job_title": job.get('title'),
            "total_matches": len(scores),
            "score_statistics": {
                "average": round(sum(scores) / len(scores), 1) if scores else 0,
                "min": round(min(scores), 1) if scores else 0,
                "max": round(max(scores), 1) if scores else 0,
                "median": round(sorted(scores)[len(scores)//2], 1) if scores else 0
            },
            "top_candidates": [
                {
                    "name": database.get_candidate(m.candidate_id).get('name'),
                    "score": m.overall_score
                } for m in matches[:5]
            ]
        }
    
    except Exception as e:
        logger.error(f"Error getting match summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# BATCH OPERATIONS
# ============================================================================

@app.post("/batch/match-all")
async def batch_match_all():
    """
    Match all candidates against all jobs
    
    Returns summary of all matches
    """
    jobs = database.get_all_jobs()
    candidates = database.get_all_candidates()
    
    if not jobs or not candidates:
        raise HTTPException(
            status_code=400,
            detail="Need at least one job and one candidate"
        )
    
    results = {}
    
    try:
        for job in jobs:
            job_id = job.get('id')
            matches = matching_engine.match_candidates(job_id)
            
            results[job_id] = {
                "job_title": job.get('title'),
                "candidate_count": len(matches),
                "top_candidates": [
                    {
                        "id": m.candidate_id,
                        "name": database.get_candidate(m.candidate_id).get('name'),
                        "score": m.overall_score
                    } for m in matches[:3]
                ]
            }
        
        return {
            "total_jobs": len(jobs),
            "total_candidates": len(candidates),
            "timestamp": datetime.utcnow().isoformat(),
            "matches": results
        }
    
    except Exception as e:
        logger.error(f"Error in batch matching: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# ERROR HANDLING
# ============================================================================

@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    """Handle HTTP exceptions"""
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """Handle general exceptions"""
    logger.error(f"Unexpected error: {exc}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error"},
    )


if __name__ == "__main__":
    uvicorn.run(
        "app.main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
